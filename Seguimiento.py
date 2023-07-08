import mysql.connector
from docx import Document
from matplotlib import pyplot as plt
from io import BytesIO
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

#DataBase Acces 

connectionBD = mysql.connector.connect(

    host="******",
    user="******",
    password="******",
    database="******"
)

cursor = connectionBD.cursor()

print("Deseas ingresar gastos o gestionar las estadisticas de tus gastos")

In = input("1 ingresar, 2 Gestionar " )

Cn = int(In)

if Cn == 1 :

        # data entered into the table
        
        id_gastos = input("Ingresa el id " )
        fecha = input("Ingresa la fecha actual " )
        categoria = input("Ingresa si es un gasto " )
        descripcion = input("Ingresa la descripcion del gasto " )
        cantidad = input("Ingresa la cantidad " )

        
        # Insert data 
        consulta = "INSERT INTO Gastos (id, fecha, categoria, descripción, cantidad) VALUES (%s, %s, %s, %s, %s)"
        datos = (id_gastos, fecha, categoria, descripcion, cantidad)
        cursor.execute(consulta, datos)

        connectionBD.commit()

        connectionBD.close()

        print("Los datos fueron ingresados correctamente")


elif Cn == 2:
        
     #Name table
    TablaO = "******"

    cursor.execute(f"SELECT COUNT(*) FROM {TablaO}")
    num_rows = cursor.fetchone()[0]

    cursor.execute(f"SELECT COUNT(*) FROM information_schema.columns WHERE table_name = '{TablaO}'")
    num_columns = cursor.fetchone()[0]

    cursor.execute(f"SELECT * FROM {TablaO}")
    filas = cursor.fetchall()

    # Create document word
    
    document = Document()
    Table = document.add_table(rows=num_rows+1, cols=num_columns)

    cursor.execute(f"SELECT COLUMN_NAME FROM information_schema.columns Where table_name = '{TablaO}'")
    names_colmns = cursor.fetchall()
    encabezados = [nombre[0] for nombre in names_colmns]

    header_celdas = Table.rows[0].cells

    for i in range(num_columns):
         header_celdas[i].text = encabezados[i]

    # agrege data on table
    for i, filas in enumerate(filas):
         row_celdas = Table.rows[i+1].cells
         for j in range(num_columns):
              row_celdas[j].text = str(filas[j])


    cursor.execute("SELECT fecha, cantidad FROM Gastos")
    resultados = cursor.fetchall()

    fechas = []
    cantidad = []

    for fila in resultados:
         fechas.append(fila[0])
         cantidad.append(fila[1])
         
    

    plt.title("GASTOS")
    plt.ylabel("Cantidad $")
    plt.xlabel("Fechas")
    plt.plot(fechas, cantidad)
    
    plt.gcf().autofmt_xdate()


    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)

    # Insertar la imagen en el documento de Word
    document.add_picture(buffer)

    guardar = input("guarda el documento con el nombre: " )
    
    # Guardar el documento
    document.save(f"{guardar}.docx")
    print("El documento se ha generado exitosamente.")


else:        
    
    print("error")
